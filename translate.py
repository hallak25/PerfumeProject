from docx import Document

from deep_translator import GoogleTranslator



# Load the original Macedonian Word document

input_path = "C: /xxx.docx"  # replace with your input and file name

output_path = "C:/xxx_translated.docx" # replace with your output path and file name





# Initialize translator and documents

translator = GoogleTranslator(source='mk', target='en')

doc_mk = Document(input_path)

doc_en = Document()



# Add heading

doc_en.add_heading("HEADER", level=1)

doc_en.add_paragraph("PARAGRAPH\n")



# Translate paragraph by paragraph

for para in doc_mk.paragraphs:

    text = para.text.strip()

    if text:

        try:

            translated = translator.translate(text)

            doc_en.add_paragraph(translated)

        except Exception as e:

            print(f"Failed to translate paragraph: {text[:40]}...")

            doc_en.add_paragraph("[Translation failed]")



# Save the translated document

doc_en.save(output_path)

print(f" Translated document saved as: {output_path}")
